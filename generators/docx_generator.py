import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0C0D0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x5C, 0x35, 0xC9)
    add_horizontal_rule(doc)


def add_body_bullet(doc, text: str):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x3F)


def add_body_text(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x3F)


def generate_docx(data: dict) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── NAME ──────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data["full_name"])
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(26)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── TITLE ─────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(data["professional_title"])
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(13)
    title_run.font.color.rgb = RGBColor(0x5C, 0x35, 0xC9)
    title_run.font.bold = True

    # ── CONTACT LINE ──────────────────────────────────────────────────────
    contact_parts = []
    if data.get("email"):       contact_parts.append(data["email"])
    if data.get("phone"):       contact_parts.append(data["phone"])
    if data.get("linkedin"):    contact_parts.append(data["linkedin"])
    if data.get("github"):      contact_parts.append(data["github"])
    if data.get("personal_website"): contact_parts.append(data["personal_website"])

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run("  •  ".join(contact_parts))
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x75)

    # ── SUMMARY ───────────────────────────────────────────────────────────
    add_section_heading(doc, "Professional Summary")
    add_body_text(doc, data.get("summary", ""))

    # ── SKILLS ────────────────────────────────────────────────────────────
    if data.get("skills"):
        add_section_heading(doc, "Skills")
        skill_chunks = [data["skills"][i:i+4] for i in range(0, len(data["skills"]), 4)]
        for chunk in skill_chunks:
            add_body_bullet(doc, "  •  ".join(chunk))

    # ── EXPERIENCE ────────────────────────────────────────────────────────
    if data.get("experience"):
        add_section_heading(doc, "Work Experience")
        for exp in data["experience"]:
            add_body_bullet(doc, exp)

    # ── PROJECTS ──────────────────────────────────────────────────────────
    if data.get("projects"):
        add_section_heading(doc, "Projects")
        for proj in data["projects"]:
            add_body_bullet(doc, proj)

    # ── EDUCATION ─────────────────────────────────────────────────────────
    if data.get("education"):
        add_section_heading(doc, "Education")
        add_body_text(doc, data["education"])

    # ── ACHIEVEMENTS ──────────────────────────────────────────────────────
    if data.get("achievements"):
        add_section_heading(doc, "Achievements")
        for line in data["achievements"].split("\n"):
            if line.strip():
                add_body_bullet(doc, line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
